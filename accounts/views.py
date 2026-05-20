
# =========================================================
# IMPORTS
# =========================================================
import os
import re
import io
import json
import subprocess
import datetime
import calendar
import requests

from groq import Groq
from decimal import Decimal

from django.shortcuts import render, redirect, get_object_or_404
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_http_methods
from django.db.models import Sum
from django.db import transaction
from django.urls import reverse
from django.utils import timezone
from datetime import timedelta
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.utils.dateparse import parse_date
from django.conf import settings
from datetime import date

from .models import Customer, MilkEntry, Payment, PRICE_PER_LITRE
from .forms import MilkEntryForm, CustomerForm, PaymentForm
from .pdf_generation import generate_bill_pdf


client = Groq(api_key=settings.GROQ_API_KEY)


# =========================================================
# DASHBOARD
# =========================================================

@login_required(login_url='login')
def home(request):
    try:
        total_customers = Customer.objects.count()
        total_ml = MilkEntry.objects.aggregate(total=Sum('quantity_ml'))['total'] or 0
        total_litres = round(Decimal(total_ml) / Decimal(1000), 2) if total_ml else Decimal(0)
        total_amount = round((Decimal(total_ml) / Decimal(1000)) * Decimal(PRICE_PER_LITRE), 2) if total_ml else Decimal(0)
        total_balance = Customer.objects.aggregate(balance=Sum('balance_amount'))['balance'] or Decimal(0)
        last_entries = MilkEntry.objects.select_related('customer').order_by('-date')[:10]

        today = timezone.localdate()
        snoozed = request.session.get('snoozed_overdue', [])
        overdue_customers = []

        for customer in Customer.objects.filter(balance_amount__gt=0):
            if customer.id in snoozed:  # type: ignore
                continue
            first_entry = (
                MilkEntry.objects
                .filter(customer=customer)
                .order_by('date')
                .first()
            )
            if not first_entry:
                continue
            age_days = (today - first_entry.date).days
            if age_days >= 90:
                severity = 'critical'
            elif age_days >= 60:
                severity = 'high'
            elif age_days >= 30:
                severity = 'medium'
            else:
                severity = 'low'
            overdue_customers.append({
                'customer':   customer,
                'balance':    round(customer.balance_amount, 2),
                'since_date': first_entry.date,
                'age_days':   age_days,
                'severity':   severity,
                'progress':   min(round((age_days / 90) * 100), 100),
            })

        overdue_customers.sort(key=lambda x: x['age_days'], reverse=True)

        context = {
            'total_customers':   total_customers,
            'total_litres':      total_litres,
            'total_balance':     round(total_balance, 2),
            'total_amount':      total_amount,
            'last_entries':      last_entries,
            'overdue_customers': overdue_customers,
            'overdue_count':     len(overdue_customers),
        }
        return render(request, 'accounts/home.html', context)
    except Exception as e:
        return render(request, 'accounts/home.html', {
            'total_customers':   0,
            'total_litres':      0,
            'total_balance':     0,
            'total_amount':      0,
            'last_entries':      [],
            'overdue_customers': [],
            'overdue_count':     0,
            'error':             str(e),
        })


@login_required(login_url='login')
def customer_list(request):
    customers = Customer.objects.all()
    for customer in customers:
        total_ml = MilkEntry.objects.filter(customer=customer).aggregate(total=Sum('quantity_ml'))['total'] or 0
        customer.total_ml = total_ml  # type: ignore
        customer.total_litres = round(Decimal(total_ml) / Decimal(1000), 2) if total_ml else Decimal(0)  # type: ignore
    return render(request, 'accounts/customer_list.html', {'customers': customers})


@login_required(login_url='login')
def customer_detail(request, customer_id):
    customer = get_object_or_404(Customer, id=customer_id)
    entries = MilkEntry.objects.filter(customer=customer).order_by('-date')

    months_data = {}
    for entry in entries:
        key = (entry.date.year, entry.date.month)
        if key not in months_data:
            months_data[key] = {
                'year':         entry.date.year,
                'month':        entry.date.month,
                'month_name':   entry.date.strftime('%B %Y'),
                'entries':      [],
                'total_ml':     0,
                'total_litres': Decimal(0),
                'total_amount': Decimal(0),
            }
        months_data[key]['entries'].append(entry)
        months_data[key]['total_ml'] += entry.quantity_ml
        months_data[key]['total_amount'] += entry.amount

    for month_key in months_data:
        ml = months_data[month_key]['total_ml']
        months_data[month_key]['total_litres'] = round(Decimal(ml) / Decimal(1000), 2)
        months_data[month_key]['total_amount'] = round(months_data[month_key]['total_amount'], 2)

    sorted_months = sorted(months_data.values(), key=lambda x: (x['year'], x['month']), reverse=True)

    context = {
        'customer':     customer,
        'months_data':  sorted_months,
        'total_entries': entries.count(),
    }
    return render(request, 'accounts/customer_detail.html', context)


@login_required(login_url='login')
@require_http_methods(["GET", "POST"])
def add_entry(request):
    if request.method == 'POST':
        form = MilkEntryForm(request.POST)
        if form.is_valid():
            customer = form.cleaned_data.get('customer')
            new_name = form.cleaned_data.get('customer_name')
            if not customer and new_name:
                customer, created = Customer.objects.get_or_create(name=new_name.strip())
            if customer:
                MilkEntry.objects.create(
                    customer=customer,
                    date=form.cleaned_data['date'],
                    quantity_ml=form.cleaned_data['quantity_ml']
                )
                return redirect('accounts:customer_list')
    else:
        form = MilkEntryForm()
    return render(request, 'accounts/entry_form.html', {'form': form})


@login_required(login_url='login')
@require_http_methods(["GET", "POST"])
def edit_entry(request, entry_id):
    entry = get_object_or_404(MilkEntry, id=entry_id)
    if request.method == 'POST':
        form = MilkEntryForm(request.POST, instance=entry)
        if form.is_valid():
            form.save()
            return redirect('accounts:customer_detail', customer_id=entry.customer.id)  # type: ignore
    else:
        form = MilkEntryForm(instance=entry)
    return render(request, 'accounts/entry_form.html', {'form': form, 'title': 'Edit Milk Entry'})


@login_required(login_url='login')
@require_http_methods(["POST"])
def delete_entry(request, entry_id):
    entry = get_object_or_404(MilkEntry, id=entry_id)
    customer_id = entry.customer.id  # type: ignore
    entry.delete()
    return redirect('accounts:customer_detail', customer_id=customer_id)


@login_required(login_url='login')
@require_http_methods(["GET", "POST"])
def edit_customer(request, customer_id):
    customer = get_object_or_404(Customer, id=customer_id)
    if request.method == 'POST':
        form = CustomerForm(request.POST, instance=customer)
        if form.is_valid():
            form.save()
            return redirect('accounts:customer_detail', customer_id=customer.id)  # type: ignore
    else:
        form = CustomerForm(instance=customer)
    return render(request, 'accounts/customer_form.html', {'form': form, 'customer': customer, 'title': 'Edit Customer'})


@login_required(login_url='login')
@require_http_methods(["POST"])
def delete_customer(request, customer_id):
    customer = get_object_or_404(Customer, id=customer_id)
    customer.delete()
    return redirect('accounts:customer_list')


@login_required(login_url='login')
def chart_data(request, customer_id):
    customer = get_object_or_404(Customer, id=customer_id)
    entries = MilkEntry.objects.filter(customer=customer).order_by('date')[:30]
    labels = [e.date.strftime('%Y-%m-%d') for e in entries]
    data = [float(e.litres) for e in entries]
    return JsonResponse({'labels': labels, 'data': data})


@login_required(login_url='login')
def bill_pdf(request, customer_id, year=None, month=None):
    customer = get_object_or_404(Customer, id=customer_id)
    if year and month:
        entries = MilkEntry.objects.filter(
            customer=customer,
            date__year=year,
            date__month=month
        ).order_by('date')
        filename = f"bill_{customer.name.replace(' ', '_')}_{year}_{month:02d}"  # type: ignore
    else:
        entries = MilkEntry.objects.filter(customer=customer).order_by('date')
        filename = f"bill_{customer.name.replace(' ', '_')}_all"  # type: ignore

    total_ml = entries.aggregate(total=Sum('quantity_ml'))['total'] or 0
    total_litres = round(Decimal(total_ml) / Decimal(1000), 2) if total_ml else Decimal(0)
    total_amount = round((Decimal(total_ml) / Decimal(1000)) * Decimal(PRICE_PER_LITRE), 2) if total_ml else Decimal(0)

    pdf_buffer = generate_bill_pdf(
        customer=customer,
        entries=entries,
        total_ml=total_ml,
        total_litres=total_litres,
        total_amount=total_amount,
        price_per_litre=PRICE_PER_LITRE,
        year=year,
        month=month
    )

    response = HttpResponse(pdf_buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
    return response


@login_required(login_url='login')
def monthly_summary(request):
    today = timezone.localdate()

    try:
        selected_month = int(request.GET.get('month', today.month))
        selected_year  = int(request.GET.get('year',  today.year))
        if not (1 <= selected_month <= 12):
            selected_month = today.month
        if selected_year < 2000 or selected_year > today.year + 1:
            selected_year = today.year
    except (ValueError, TypeError):
        selected_month = today.month
        selected_year  = today.year

    start_current = date(selected_year, selected_month, 1)
    last_day = calendar.monthrange(selected_year, selected_month)[1]
    end_current = date(selected_year, selected_month, last_day)

    if selected_month == 1:
        prev_year, prev_month = selected_year - 1, 12
    else:
        prev_year, prev_month = selected_year, selected_month - 1

    start_previous = date(prev_year, prev_month, 1)
    last_day_prev  = calendar.monthrange(prev_year, prev_month)[1]
    end_previous   = date(prev_year, prev_month, last_day_prev)

    def build_summary(start_date, end_date):
        entries = MilkEntry.objects.filter(
            date__range=[start_date, end_date]
        ).select_related('customer')

        summary = {}
        total_ml_all = 0
        for entry in entries:
            cid = entry.customer.id  # type: ignore
            if cid not in summary:
                summary[cid] = {
                    'name':     entry.customer.name,
                    'total_ml': 0,
                    'amount':   Decimal(0),
                }
            summary[cid]['total_ml'] += entry.quantity_ml
            summary[cid]['amount']   += entry.amount
            total_ml_all             += entry.quantity_ml

        summary_list = []
        for data in summary.values():
            litres = round(Decimal(data['total_ml']) / Decimal(1000), 2)
            amt    = round(data['amount'], 2)
            summary_list.append({
                'name':     data['name'],
                'total_ml': data['total_ml'],
                'litres':   litres,
                'amount':   amt,
            })

        total_amount = round(sum(item['amount'] for item in summary_list), 2)
        total_litres = round(Decimal(total_ml_all) / Decimal(1000), 2)
        return summary_list, total_amount, total_litres

    current_summary,  current_total_amount,  current_total_litres  = build_summary(start_current,  end_current)
    previous_summary, previous_total_amount, previous_total_litres = build_summary(start_previous, end_previous)

    month_names = [
        'January','February','March','April','May','June',
        'July','August','September','October','November','December'
    ]
    months_list = [{'value': i + 1, 'label': month_names[i]} for i in range(12)]

    earliest = MilkEntry.objects.order_by('date').values_list('date', flat=True).first()
    earliest_year = earliest.year if earliest else today.year
    years_list = list(range(today.year, earliest_year - 1, -1))

    current_tab_label  = f"{month_names[selected_month - 1]} {selected_year}"
    previous_tab_label = f"{month_names[prev_month - 1]} {prev_year}"

    return render(request, 'accounts/monthly_summary.html', {
        'selected_month':  selected_month,
        'selected_year':   selected_year,
        'months_list':     months_list,
        'years_list':      years_list,
        'current_tab_label':  current_tab_label,
        'previous_tab_label': previous_tab_label,
        'current_summary':       current_summary,
        'current_total_amount':  current_total_amount,
        'current_total_litres':  current_total_litres,
        'start_current':         start_current,
        'end_current':           end_current,
        'previous_summary':       previous_summary,
        'previous_total_amount':  previous_total_amount,
        'previous_total_litres':  previous_total_litres,
        'start_previous':         start_previous,
        'end_previous':           end_previous,
    })


# =========================================================
# AI DOCUMENT PROCESSING -- UPLOAD VIEWS
# =========================================================

@login_required(login_url='login')
def upload_entries(request):
    if request.method == "POST":
        uploaded_file = request.FILES.get("file")
        if not uploaded_file:
            return render(request, "accounts/upload_entries.html", {"error": "No file selected"})

        path = default_storage.save(
            f"uploads/{uploaded_file.name}",
            ContentFile(uploaded_file.read())
        )

        try:
            extracted_data = process_file_with_ai(path)
        except Exception as e:
            return render(request, "accounts/upload_entries.html", {
                "error": f"AI Processing Failed: {str(e)}"
            })

        request.session['ai_preview'] = extracted_data
        return redirect("accounts:preview_entries")

    return render(request, "accounts/upload_entries.html")


@login_required(login_url='login')
def preview_entries(request):
    preview_data = request.session.get('ai_preview', [])
    customers = Customer.objects.all()

    existing_entries = list(MilkEntry.objects.values_list("customer__name", "date"))
    existing_entries = [[name, str(d)] for name, d in existing_entries]

    # ensure_ascii=False keeps Tamil chars intact; |safe used in template (we control this data)
    existing_entries_json = json.dumps(existing_entries, ensure_ascii=False)

    return render(request, "accounts/preview_entries.html", {
        "entries":          preview_data,
        "customers":        customers,
        "existing_entries": existing_entries_json,
    })


@login_required(login_url='login')
@require_http_methods(["POST"])
def confirm_entries(request):
    total_rows = int(request.POST.get("total_rows", 0))
    saved_entries = []
    skipped_entries = []

    with transaction.atomic():
        for i in range(total_rows):
            try:
                if not request.POST.get(f"include_{i}"):
                    continue

                customer_name = request.POST.get(f"customer_{i}", "").strip()
                date_str      = request.POST.get(f"date_{i}")
                quantity_raw  = request.POST.get(f"quantity_{i}", "0")

                if not customer_name or not date_str:
                    skipped_entries.append({"row": i, "reason": "Missing data"})
                    continue

                parsed_date = parse_date(date_str)
                if not parsed_date:
                    skipped_entries.append({"row": i, "reason": "Invalid date"})
                    continue

                try:
                    quantity = int(quantity_raw)
                except ValueError:
                    skipped_entries.append({"row": i, "reason": "Invalid quantity"})
                    continue

                if quantity < 0 or quantity > 10000:
                    skipped_entries.append({"row": i, "reason": "Suspicious quantity"})
                    continue

                customer = (
                    Customer.objects.filter(name__iexact=customer_name).first()
                    or Customer.objects.create(name=customer_name)
                )

                if MilkEntry.objects.filter(customer=customer, date=parsed_date).exists():
                    skipped_entries.append({"row": i, "reason": "Duplicate entry"})
                    continue

                MilkEntry.objects.create(
                    customer=customer,
                    date=parsed_date,
                    quantity_ml=quantity
                )

                saved_entries.append({
                    "customer":    customer.name,
                    "date":        str(parsed_date),
                    "quantity_ml": quantity,
                })

            except Exception as e:
                skipped_entries.append({"row": i, "reason": str(e)})
                continue

    audit_path = settings.BASE_DIR / "ai_audit_log.txt"
    audit_record = {
        "user":            request.user.username,
        "saved":           saved_entries,
        "skipped":         skipped_entries,
        "total_processed": total_rows,
    }
    with open(audit_path, "a", encoding="utf-8") as log:
        log.write(json.dumps(audit_record) + "\n")

    request.session.pop("ai_preview", None)
    return redirect("accounts:customer_list")


# =========================================================
# AI CORE -- FILE EXTRACTION ENGINE
# =========================================================

def _call_ocrspace_api(file_bytes, filename, language="eng"):
    """
    Send raw bytes to the OCR.Space API and return the extracted text.

    Supported language codes:
      "eng" - English  |  "tam" - Tamil
    OCR.Space free tier: 25,000 requests/month, max 1 MB per file.
    Set OCRSPACE_API_KEY in Django settings (free demo key: "helloworld").

    OCREngine 2 is used for superior handwriting recognition.
    isTable=True preserves column/row structure (name + quantity columns).
    """
    api_key = getattr(settings, "OCRSPACE_API_KEY", "helloworld")

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "jpg"
    mime_map = {
        "jpg": "image/jpeg", "jpeg": "image/jpeg",
        "png": "image/png",  "bmp": "image/bmp",
        "tiff": "image/tiff","tif": "image/tiff",
        "gif": "image/gif",  "webp": "image/webp",
        "pdf": "application/pdf",
    }
    mime = mime_map.get(ext, "image/jpeg")

    payload = {
        "apikey":             api_key,
        "language":           language,
        "isOverlayRequired":  "false",
        "detectOrientation":  "true",
        "scale":              "true",   # auto-upscale small images
        "OCREngine":          "2",      # Engine 2 = better for handwriting
        "isTable":            "true",   # preserve column structure
    }

    response = requests.post(
        "https://api.ocr.space/parse/image",
        data=payload,
        files={"file": (filename, file_bytes, mime)},
        timeout=60,
    )
    response.raise_for_status()

    result = response.json()
    if result.get("IsErroredOnProcessing"):
        error_msg = result.get("ErrorMessage", ["Unknown OCR.Space error"])
        raise Exception(f"OCR.Space error: {error_msg}")

    parsed_results = result.get("ParsedResults") or []
    text = "\n".join(r.get("ParsedText", "") for r in parsed_results).strip()

    print("\n========== OCR.Space TEXT ==========")
    print(text[:500])
    print("=====================================\n")
    return text


def _extract_from_image(file_path):
    """
    Extract text from an image using the OCR.Space API.
    Makes two passes (English + Tamil) and merges results so that
    mixed-script dairy registers are handled correctly.
    """
    filename = os.path.basename(file_path)

    with default_storage.open(file_path, "rb") as f:
        file_bytes = f.read()

    # Pass 1: English - digits, column headers, date labels
    text_eng = _call_ocrspace_api(file_bytes, filename, language="eng")

    # Pass 2: Tamil - customer names written in Tamil script
    try:
        text_tam = _call_ocrspace_api(file_bytes, filename, language="tam")
    except Exception as exc:
        print(f"[OCR] Tamil pass failed (non-fatal): {exc}")
        text_tam = ""

    if len(text_tam) > len(text_eng):
        return (text_tam + "\n" + text_eng).strip()
    return (text_eng + "\n" + text_tam).strip()


def _extract_from_pdf(file_path):
    """
    Try pdfplumber for text-based PDFs.
    Fall back to OCR.Space for scanned / image-only PDFs.
    """
    try:
        import pdfplumber  # pip install pdfplumber
    except ImportError:
        raise Exception("pdfplumber is required for PDF extraction. Run: pip install pdfplumber")

    with default_storage.open(file_path, "rb") as f:
        pdf_bytes = f.read()

    def _reconstruct_lines_from_words(page):
        words = page.extract_words(keep_blank_chars=False, x_tolerance=5, y_tolerance=5)
        if not words:
            return ""
        rows = {}
        for w in words:
            y_key = round(w["top"] / 3) * 3
            rows.setdefault(y_key, []).append(w)
        lines = []
        for y_key in sorted(rows.keys()):
            row_words = sorted(rows[y_key], key=lambda w: w["x0"])
            lines.append("  ".join(w["text"] for w in row_words))
        return "\n".join(lines)

    pages_text = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text(layout=True) or ""
            words_text = _reconstruct_lines_from_words(page)
            best = words_text if len(words_text) >= len(text) else text
            pages_text.append(best)
            print(f"[PDF] Page extracted ({len(best)} chars):\n{best[:300]}")

    text = "\n".join(pages_text).strip()

    if text:
        return text

    # Scanned PDF fallback: send each page image to OCR.Space
    print("[PDF] No selectable text -- falling back to OCR.Space per page")
    try:
        from pdf2image import convert_from_bytes  # pip install pdf2image + poppler
    except ImportError:
        raise Exception(
            "pdf2image is required for scanned PDF OCR. "
            "Run: pip install pdf2image  (also needs poppler: "
            "apt install poppler-utils  /  winget install poppler)"
        )

    images = convert_from_bytes(pdf_bytes, dpi=200)
    all_text_parts = []
    for page_num, pil_img in enumerate(images, start=1):
        buf = io.BytesIO()
        pil_img.save(buf, format="PNG")
        img_bytes = buf.getvalue()
        page_filename = f"page_{page_num}.png"

        text_eng = _call_ocrspace_api(img_bytes, page_filename, language="eng")
        try:
            text_tam = _call_ocrspace_api(img_bytes, page_filename, language="tam")
        except Exception as exc:
            print(f"[PDF OCR] Tamil pass page {page_num} failed (non-fatal): {exc}")
            text_tam = ""

        if len(text_tam) > len(text_eng):
            all_text_parts.append((text_tam + "\n" + text_eng).strip())
        else:
            all_text_parts.append((text_eng + "\n" + text_tam).strip())

    return "\n".join(all_text_parts)


def _extract_from_docx(file_path):
    """Extract text from .docx using python-docx."""
    try:
        from docx import Document  # pip install python-docx
    except ImportError:
        raise Exception("python-docx is required. Run: pip install python-docx")

    with default_storage.open(file_path, "rb") as f:
        doc = Document(io.BytesIO(f.read()))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def _extract_from_doc(file_path):
    """Legacy .doc via antiword."""
    with default_storage.open(file_path, "rb") as f:
        doc_bytes = f.read()
    tmp = os.path.join(os.environ.get("TEMP", "/tmp"), "_uploaded.doc")
    with open(tmp, "wb") as tf:
        tf.write(doc_bytes)
    try:
        result = subprocess.run(["antiword", tmp], capture_output=True, text=True)
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except FileNotFoundError:
        pass
    raise Exception(
        "Could not read .doc file. Install antiword: "
        "https://www.winfield.demon.nl/ (Windows) or apt install antiword (Linux)"
    )


def _extract_from_xlsx(file_path):
    """Extract text from .xlsx / .xls / .ods spreadsheets."""
    with default_storage.open(file_path, "rb") as f:
        raw = f.read()
    try:
        import openpyxl  # pip install openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
        rows = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                line = "\t".join(str(c) if c is not None else "" for c in row)
                if line.strip():
                    rows.append(line)
        return "\n".join(rows)
    except Exception:
        import pandas as pd
        df = pd.read_excel(io.BytesIO(raw))
        return df.to_string(index=False)


def _extract_from_csv(file_path):
    with default_storage.open(file_path, "rb") as f:
        return f.read().decode("utf-8", errors="ignore")


def _extract_from_txt(file_path):
    with default_storage.open(file_path, "rb") as f:
        return f.read().decode("utf-8", errors="ignore")


def _extract_from_rtf(file_path):
    with default_storage.open(file_path, "rb") as f:
        raw = f.read().decode("utf-8", errors="ignore")
    text = re.sub(r'\{[^{}]*\}', '', raw)
    text = re.sub(r'\\[a-z]+\d* ?', '', text)
    text = re.sub(r'[{}\\]', '', text)
    return text.strip()


def _extract_from_odt(file_path):
    try:
        from odf import teletype  # type: ignore  # pip install odfpy
        from odf.opendocument import load as odf_load  # type: ignore
    except ImportError:
        raise Exception("odfpy is required for .odt extraction. Run: pip install odfpy")

    with default_storage.open(file_path, "rb") as f:
        doc = odf_load(io.BytesIO(f.read()))
    return teletype.extractText(doc.text)  # type: ignore


# =========================================================
# DISPATCH TABLE
# =========================================================

IMAGE_EXTS = {"jpg", "jpeg", "png", "bmp", "tiff", "tif", "webp", "gif"}

EXTRACTOR_MAP = {
    **{ext: _extract_from_image for ext in IMAGE_EXTS},
    "pdf":  _extract_from_pdf,
    "docx": _extract_from_docx,
    "doc":  _extract_from_doc,
    "odt":  _extract_from_odt,
    "rtf":  _extract_from_rtf,
    "xlsx": _extract_from_xlsx,
    "xls":  _extract_from_xlsx,
    "ods":  _extract_from_xlsx,
    "csv":  _extract_from_csv,
    "tsv":  _extract_from_csv,
    "txt":  _extract_from_txt,
    "md":   _extract_from_txt,
    "log":  _extract_from_txt,
    "text": _extract_from_txt,
}


# =========================================================
# MAIN AI PROCESSING FUNCTION
# =========================================================

def _extract_json_array(text):
    """
    Extract the first complete JSON array from text using bracket counting.
    Handles greedy-regex pitfalls: trailing notes, nested structures,
    escaped quotes inside strings.
    Returns a parsed list, or None if nothing valid is found.
    """
    start = text.find('[')
    if start == -1:
        return None

    depth = 0
    in_string = False
    escape_next = False

    for i, ch in enumerate(text[start:], start):
        if escape_next:
            escape_next = False
            continue
        if ch == '\\' and in_string:
            escape_next = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == '[':
            depth += 1
        elif ch == ']':
            depth -= 1
            if depth == 0:
                candidate = text[start:i + 1]
                try:
                    return json.loads(candidate)
                except json.JSONDecodeError:
                    return None

    return None


def process_file_with_ai(file_path):
    extension = file_path.rsplit(".", 1)[-1].lower()
    extractor = EXTRACTOR_MAP.get(extension)

    if extractor is None:
        try:
            with default_storage.open(file_path, "rb") as f:
                text_content = f.read().decode("utf-8", errors="ignore")
            if not text_content.strip():
                raise ValueError("empty")
        except Exception:
            raise Exception(
                f"Unsupported file type: .{extension}. "
                "Supported: " + ", ".join(sorted(EXTRACTOR_MAP.keys()))
            )
    else:
        text_content = extractor(file_path)

    if not text_content.strip():
        raise Exception("Could not extract readable text from the document.")

    groq_client = Groq(api_key=settings.GROQ_API_KEY)

    prompt = f"""
You are an intelligent dairy billing assistant. Extract milk supply entries from the text below.

Rules:
- Customer names may be in English OR Tamil script - preserve them exactly as written.
- IMPORTANT: Customer names in a dairy register are often NICKNAMES or UNUSUAL WORDS.
  A word like "laptop", "scooter", "mobile", "tower" etc. that appears on the LEFT side
  of a row followed by a number IS the customer nickname. Do NOT skip it, do NOT replace
  it with null, and do NOT rename it. Extract it exactly as written.
- The document structure is a list where each row is:
      <customer_name>   <quantity_in_ml>
  Every row with a name on the left AND a number on the right is one entry.
  Do NOT skip any row. Do NOT leave the customer field blank or null.
- Dates may appear in many formats. Normalise ALL dates to ISO format YYYY-MM-DD:
    * "February 2026" or "Feb 2026"  ->  "2026-02-01"
    * "01/02/2026" or "1-2-26"       ->  "2026-02-01"
    * "01 Feb 26"                    ->  "2026-02-01"
    * "14-may-2026"                  ->  "2026-05-14"
    * A bare day number like "5" inside a document already headed "March 2026" -> "2026-03-05"
- A date header applies to ALL rows below it until the next date header appears.
- Extract quantity in ml. If a litre value is given (e.g. "2.5 L"), convert to ml (2500).
- If only a bare number exists with no unit, assume ml.
- Column headers like "Ml" or "ML" are NOT customer names - skip them.
- If a 0 ml quantity is detected, accept it (zero supply that day is still a valid entry).

CRITICAL output format:
- Your ENTIRE response must be a single raw JSON array.
- No explanations, no markdown, no code fences, no backticks, no preamble.
- If no entries are found return exactly: []
- Do NOT wrap the array in an object. Start with [ and end with ].
- Each element: {{"customer": "<name>", "date": "YYYY-MM-DD", "quantity_ml": <integer>}}

Text:
{text_content}
"""

    response = groq_client.chat.completions.create(
        model=settings.GROQ_MODEL,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a JSON-only data extractor. "
                    "You NEVER use markdown, code fences, or explanations. "
                    "You output ONLY raw JSON arrays, nothing else. "
                    "Customer names may contain Tamil or other non-ASCII characters - preserve them exactly."
                )
            },
            {"role": "user", "content": prompt}
        ],
        temperature=0
    )

    raw_output = response.choices[0].message.content.strip()  # type: ignore

    print("[AI RAW OUTPUT]", raw_output[:500])

    clean_output = re.sub(r"```(?:json)?\s*", "", raw_output).strip()
    clean_output = re.sub(r"```\s*$", "", clean_output, flags=re.MULTILINE).strip()

    parsed = _extract_json_array(clean_output)

    if parsed is None:
        obj_match = re.search(r"\{.*?\}", clean_output, re.DOTALL)
        if obj_match:
            try:
                parsed = [json.loads(obj_match.group(0))]
            except json.JSONDecodeError:
                parsed = None

    if parsed is None:
        raise Exception(
            f"AI did not return valid JSON. "
            f"Raw response (first 300 chars): {raw_output[:300]!r}"
        )

    preview_list = []
    for row in parsed:
        try:
            quantity = int(row.get("quantity_ml", 0))
        except Exception:
            quantity = 0

        confidence = 0.2 if (quantity < 0 or quantity > 50000) else 0.9

        preview_list.append({
            "customer":    row.get("customer"),
            "date":        row.get("date"),
            "quantity_ml": quantity,
            "confidence":  confidence,
            "valid":       confidence > 0.5,
        })

    return preview_list


# =========================================================
# PAYMENT VIEWS
# =========================================================

def add_payment(request, customer_id):
    customer = get_object_or_404(Customer, id=customer_id)

    today = timezone.localdate()
    first_of_this_month = today.replace(day=1)
    last_month_end   = first_of_this_month - datetime.timedelta(days=1)
    last_month_start = last_month_end.replace(day=1)

    total_qty_ml = (
        MilkEntry.objects
        .filter(
            customer=customer,
            date__gte=last_month_start,
            date__lte=last_month_end,
        )
        .aggregate(total=Sum('quantity_ml'))['total'] or 0
    )
    prev_month_total = (
        Decimal(total_qty_ml) / Decimal(1000) * Decimal(str(PRICE_PER_LITRE))
    )

    total_outstanding = customer.balance_amount + prev_month_total

    if request.method == 'POST':
        form = PaymentForm(request.POST)
        if form.is_valid():
            paid_amount = form.cleaned_data['amount']

            if paid_amount <= 0:
                form.add_error('amount', 'Payment amount must be greater than zero.')
            else:
                payment = form.save(commit=False)
                payment.customer = customer
                payment.save()

                customer.balance_amount = total_outstanding - paid_amount
                customer.save()

                messages.success(
                    request,
                    f'Payment of Rs.{paid_amount:.2f} recorded. '
                    f'New balance: Rs.{customer.balance_amount:.2f}'
                )
                return redirect('accounts:customer_detail', customer_id=customer.id)  # type: ignore
    else:
        form = PaymentForm()

    recent_payments = (
        Payment.objects
        .filter(customer=customer)
        .order_by('-paid_on')[:5]
    )

    return render(request, 'accounts/payment_form.html', {
        'form':              form,
        'customer':          customer,
        'prev_month_total':  prev_month_total,
        'total_outstanding': total_outstanding,
        'recent_payments':   recent_payments,
    })


# =========================================================
# MISC
# =========================================================

@login_required(login_url='login')
def payment_history(request, customer_id=None):
    selected_customer = None
    if customer_id:
        selected_customer = get_object_or_404(Customer, id=customer_id)
        payments = Payment.objects.filter(customer=selected_customer)
    else:
        payments = Payment.objects.select_related('customer').all()

    filter_customer_id = request.GET.get('customer_id')
    if filter_customer_id and not customer_id:
        try:
            selected_customer = Customer.objects.get(id=filter_customer_id)
            payments = payments.filter(customer=selected_customer)
        except Customer.DoesNotExist:
            pass

    filter_month = request.GET.get('month')
    if filter_month:
        try:
            year, month = filter_month.split('-')
            payments = payments.filter(paid_on__year=int(year), paid_on__month=int(month))
        except (ValueError, AttributeError):
            pass

    payments = payments.order_by('-paid_on', '-created_at')

    total_paid    = payments.aggregate(total=Sum('amount'))['total'] or Decimal('0.00')
    all_customers = Customer.objects.order_by('name')

    customer_summary = (
        Payment.objects
        .values('customer__id', 'customer__name')
        .annotate(total=Sum('amount'))
        .order_by('-total')
    )

    context = {
        'payments':          payments,
        'total_paid':        round(total_paid, 2),
        'selected_customer': selected_customer,
        'all_customers':     all_customers,
        'customer_summary':  customer_summary,
        'filter_month':      filter_month or '',
        'payment_count':     payments.count(),
    }
    return render(request, 'accounts/payment_history.html', context)


@login_required(login_url='login')
def edit_payment(request, payment_id):
    payment  = get_object_or_404(Payment, id=payment_id)
    customer = payment.customer

    if request.method == 'POST':
        form = PaymentForm(request.POST, instance=payment)
        if form.is_valid():
            old_amount = payment.amount
            new_amount = form.cleaned_data['amount']

            form.save()

            customer.balance_amount = customer.balance_amount + old_amount - new_amount
            customer.save()

            messages.success(
                request,
                f'Payment updated. New balance: Rs.{customer.balance_amount:.2f}'
            )
            return redirect('accounts:payment_history')
    else:
        form = PaymentForm(instance=payment)

    return render(request, 'accounts/edit_payment.html', {
        'form':     form,
        'payment':  payment,
        'customer': customer,
    })


@login_required(login_url='login')
@require_http_methods(["POST"])
def delete_payment(request, payment_id):
    payment  = get_object_or_404(Payment, id=payment_id)
    customer = payment.customer

    customer.balance_amount += payment.amount
    customer.save()

    payment.delete()

    messages.success(
        request,
        f'Payment of Rs.{payment.amount:.2f} deleted. '
        f'Balance restored to Rs.{customer.balance_amount:.2f}'
    )
    return redirect('accounts:payment_history')


def calculator_view(request):
    return render(request, "accounts/calculator.html")
